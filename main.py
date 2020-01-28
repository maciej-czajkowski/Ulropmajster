import tkinter as tk
import tkinter.ttk as ttk
import functions as fun
import docx
import os



drivers = fun.getdrivers()

driver_name_list = []
for i in range(len(drivers)):
    driver_name_list.append(drivers[i]['name'])

employees = fun.getemployees()

employee_name_list = []
for i in range(len(employees)):
    employee_name_list.append(employees[i]['name'])

class GUI():

    def __init__(self):
        self.root = tk.Tk()
        self.root.title("UlropMajster!")
        self.root.wm_iconbitmap('data\ico.ico')

        # menu top = left + right
        self.menu_top = tk.Frame(self.root, width = 350, height = 100, bg ='white')
        self.menu_top.pack(side = tk.TOP)

        # menu left
        self.menu_left_upper = tk.Frame(self.menu_top, width=150, height = 100, bg="white")
        self.menu_left_upper.pack(side="left")

        # menu right
        self.menu_right = tk.Frame(self.menu_top, width = 200, height = 100, bg='white')
        self.menu_right.pack(side = 'left')

        # menu bottom
        self.menu_bottom = tk.Frame(self.root, width = 350, height = 50, bg = 'white')
        self.menu_bottom.pack(side = tk.BOTTOM)

        self.dname = tk.Label(self.menu_left_upper, height = 1, text="Imię i nazwisko kierowcy", anchor = tk.W, bg='white')
        self.ename = tk.Label(self.menu_left_upper, height = 1,text="Imię i nazwisko podpisującego", anchor = tk.W, bg='white')
        self.start = tk.Label(self.menu_left_upper, height = 1,text="Od", anchor = tk.W, bg='white')
        self.end = tk.Label(self.menu_left_upper, height = 1,text="Do", anchor = tk.W, bg='white')
        self.dname.pack(fill='x')
        self.ename.pack(fill='x')
        self.start.pack(fill='x')
        self.end.pack(fill='x')



        self.dcombo = ttk.Combobox(self.menu_right, width = 50, height = 50,)
        self.dcombo['values'] = driver_name_list
        self.dcombo.pack(side = 'top')
        self.ecombo = ttk.Combobox(self.menu_right, width=50, height=50, )
        self.ecombo['values'] = employee_name_list
        self.ecombo.pack(side='top')

        self.startframe = tk.Frame(self.menu_right, bg = "black")
        self.startframe.pack(side = tk.TOP)

        self.comboh = ttk.Combobox(self.startframe, width = 2)
        self.comboh['values'] = fun.naturals(24)
        self.comboh.pack(side = tk.LEFT)
        self.labelcoma = tk.Label(self.startframe, text = ':')
        self.labelcoma.pack(side = tk.LEFT)
        self.combomin = ttk.Combobox(self.startframe, width = 2)
        self.combomin['values'] = fun.naturals(60)
        self.combomin.pack(side = tk.LEFT)
        self.labelspace = tk.Label(self.startframe, text='  ')
        self.labelspace.pack(side=tk.LEFT)
        self.combod = ttk.Combobox(self.startframe, width=2)
        self.combod['values'] = fun.naturals2(1, 32)
        self.combod.pack(side=tk.LEFT)
        self.labelsrednik = tk.Label(self.startframe, text='-')
        self.labelsrednik.pack(side=tk.LEFT)
        self.combom = ttk.Combobox(self.startframe, width=2)
        self.combom['values'] = fun.naturals2(1, 13)
        self.combom.pack(side=tk.LEFT)
        self.labelsrednik2= tk.Label(self.startframe, text='-')
        self.labelsrednik2.pack(side = tk.LEFT)
        self.comboy = ttk.Combobox(self.startframe, width=4)
        self.comboy['values'] = fun.naturals2(2019, 2031)
        self.comboy.pack(side=tk.LEFT)

        self.endframe = tk.Frame(self.menu_right, bg="black")
        self.endframe.pack(side=tk.BOTTOM)

        self.comboh2 = ttk.Combobox(self.endframe, width=2)
        self.comboh2['values'] = fun.naturals(24)
        self.comboh2.pack(side=tk.LEFT)
        self.labelcoma2 = tk.Label(self.endframe, text=':')
        self.labelcoma2.pack(side=tk.LEFT)
        self.combomin2 = ttk.Combobox(self.endframe, width=2)
        self.combomin2['values'] = fun.naturals(60)
        self.combomin2.pack(side=tk.LEFT)
        self.labelspace2 = tk.Label(self.endframe, text='  ')
        self.labelspace2.pack(side=tk.LEFT)
        self.combod2 = ttk.Combobox(self.endframe, width=2)
        self.combod2['values'] = fun.naturals2(1, 32)
        self.combod2.pack(side=tk.LEFT)
        self.labelsrednik2 = tk.Label(self.endframe, text='-')
        self.labelsrednik2.pack(side=tk.LEFT)
        self.combom2 = ttk.Combobox(self.endframe, width=2)
        self.combom2['values'] = fun.naturals2(1, 13)
        self.combom2.pack(side=tk.LEFT)
        self.labelsrednik22 = tk.Label(self.endframe, text='-')
        self.labelsrednik22.pack(side=tk.LEFT)
        self.comboy2 = ttk.Combobox(self.endframe, width=4)
        self.comboy2['values'] = fun.naturals2(2019, 2031)
        self.comboy2.pack(side=tk.LEFT)

        def click_print():
            data_to_print = []
            data_to_print.append(self.dcombo.get())
            data_to_print.append(self.ecombo.get())
            data_to_print.append(self.comboh.get())
            data_to_print.append(self.combomin.get())
            data_to_print.append(self.combod.get())
            data_to_print.append(self.combom.get())
            data_to_print.append(self.comboy.get())
            data_to_print.append(self.comboh2.get())
            data_to_print.append(self.combomin2.get())
            data_to_print.append(self.combod2.get())
            data_to_print.append(self.combom2.get())
            data_to_print.append(self.comboy2.get())
            print(data_to_print)
            doc = docx.Document('data\wzor2.docx')
            doc.tables[0].rows[0].cells[0].paragraphs[1].text = 'Nazwa przedsiębiorstwa: Trans Express Sp. z o.o. '
            doc.tables[0].rows[0].cells[0].paragraphs[7].text = 'Imię i nazwisko: ' + data_to_print[1]
            temp = 0
            for i in range(len(employees)):
                if  employees[i]['name'] == data_to_print[1]:
                    temp = i
            doc.tables[0].rows[0].cells[0].paragraphs[8].text = 'Stanowisko w przedsiębiorstwie: ' + employees[temp]['position']
            doc.tables[0].rows[0].cells[0].paragraphs[10].text = 'Imię i nazwisko: ' + data_to_print[0]
            temp = 0
            for i in range(len(employees)):
                if  employees[i]['name'] == data_to_print[0]:
                    temp = i

            doc.tables[0].rows[0].cells[0].paragraphs[11].text = 'Data urodzenia (dzień-miesiąc-rok): ' + drivers[temp]['date_of_birth']
            doc.tables[0].rows[0].cells[0].paragraphs[12].text = 'Numer prawa jazdy lub dowodu osobistego lub paszportu: ' + drivers[temp]['license']
            doc.tables[0].rows[0].cells[0].paragraphs[13].text = 'który rozpoczął pracę w przedsiębiorstwie dnia (dzień-miesiąc-rok): ' + drivers[temp]['employed']
            doc.tables[0].rows[0].cells[0].paragraphs[15].text = 'od (godzina-dzień-miesiąc-rok): ' + data_to_print[2] + ":" + data_to_print[3] + " " + data_to_print[4] + "-" + data_to_print[5] + "-" + data_to_print[6]
            doc.tables[0].rows[0].cells[0].paragraphs[16].text = 'do (godzina-dzień-miesiąc-rok): ' + data_to_print[7] + ":" + data_to_print[8] + " " + data_to_print[9] + "-" + data_to_print[10] + "-" + data_to_print[11]
            doc.tables[0].rows[0].cells[0].paragraphs[23].text = "Miejscowość: Głowno  Data: " + data_to_print[9] + "-" + data_to_print[10] + "-" + data_to_print[11]
            doc.paragraphs[6].text = "Miejscowość: Głowno  Data: " + data_to_print[9] + "-" + data_to_print[10] + "-" + data_to_print[11]

            url = os.getcwd() + "\_ulropy" + "\\" + data_to_print[0] + '\\'

            if not os.path.exists(url):
                os.makedirs(url)
            url = os.getcwd() + "\_ulropy" + "\\" + data_to_print[0] + '\\' + data_to_print[11] + data_to_print[10] +  data_to_print[9] + data_to_print[0] + ".docx"
            doc.save(url)
            print(url)
            os.startfile(url)



        self.printbutton = tk.Button(self.menu_bottom, width = 10, text = 'PRINT', bg="white", command = click_print)
        self.printbutton.pack()




        self.root.mainloop()
GUI()



